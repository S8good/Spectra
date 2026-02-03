# nanosense/utils/report_generator.py

import os
import time
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
from datetime import datetime
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from nanosense.algorithms.peak_analysis import find_main_resonance_peak, calculate_fwhm


def run_analysis_pipeline(wavelengths, spectra_df):
    """
    分析流水线，负责对加载的光谱数据进行核心计算。
    """
    try:
        peak_metrics_list = []

        # 1. 遍历每一条光谱（DataFrame的每一列）
        for col_name in spectra_df.columns:
            spectrum_data = spectra_df[col_name].values

            # 2. 寻找主共振峰
            #    这里使用最简单可靠的 np.argmax 来找到最大值点作为峰值
            peak_index, _ = find_main_resonance_peak(spectrum_data, wavelengths, min_height=-np.inf)

            if peak_index is not None:
                peak_wl = wavelengths[peak_index]
                peak_int = spectrum_data[peak_index]

                # 3. 计算半峰全宽 (FWHM)
                fwhm_results = calculate_fwhm(wavelengths, spectrum_data, [peak_index])
                fwhm = fwhm_results[0] if fwhm_results else np.nan

                peak_metrics_list.append({
                    'Spectrum Name': col_name,
                    'Peak Wavelength (nm)': peak_wl,
                    'Peak Intensity': peak_int,
                    'FWHM (nm)': fwhm
                })

        # 4. 将结果整理成 DataFrame
        summary_df = pd.DataFrame(peak_metrics_list)

        # 5. 计算统计数据 (均值, 标准差, 变异系数)
        stats_df = summary_df[['Peak Wavelength (nm)', 'Peak Intensity', 'FWHM (nm)']].agg(['mean', 'std']).T
        stats_df['CV (%)'] = (stats_df['std'] / stats_df['mean']) * 100

        # 6. 计算平均光谱
        average_spectrum_y = spectra_df.mean(axis=1).values
        average_spectrum_df = pd.DataFrame({
            'Wavelength (nm)': wavelengths,
            'Average Intensity': average_spectrum_y
        })

        # 7. 返回所有分析结果
        return {
            'summary_df': summary_df,
            'stats_df': stats_df,
            'average_spectrum_df': average_spectrum_df,
            'full_spectra_df': pd.concat([pd.DataFrame({'Wavelength (nm)': wavelengths}), spectra_df], axis=1)
        }
    except Exception as e:
        return {'error': str(e)}


def generate_pdf_report(report_subfolder, input_filename, analysis_results):
    """
    使用ReportLab生成专业的PDF报告。
    """
    try:
        pdf_path = os.path.join(report_subfolder, f"{input_filename}_report.pdf")
        doc = SimpleDocTemplate(pdf_path, pagesize=letter,
                                rightMargin=0.75*inch, leftMargin=0.75*inch,
                                topMargin=0.75*inch, bottomMargin=0.75*inch)
        story = []
        styles = getSampleStyleSheet()
        
        # 自定义标题样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f4788'),
            spaceAfter=30,
            alignment=1  # 居中
        )
        
        # 标题
        title = Paragraph("Spectrum Analysis Report", title_style)
        story.append(title)
        
        # 生成日期
        date_text = Paragraph(f"<i>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</i>", 
                             styles['Normal'])
        story.append(date_text)
        story.append(Spacer(1, 0.3*inch))
        
        # 峰值指标摘要表
        story.append(Paragraph("1. Peak Metrics Summary", styles['Heading2']))
        summary_data = [analysis_results['summary_df'].columns.tolist()] + analysis_results['summary_df'].values.tolist()
        summary_table = Table(summary_data, colWidths=[2*inch, 1.2*inch, 1.2*inch, 1.2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F2F2F2')])
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 0.3*inch))
        
        # 统计数据表
        story.append(Paragraph("2. Statistical Summary", styles['Heading2']))
        stats_data = [[str(idx)] + [f"{val:.4f}" if isinstance(val, float) else str(val) 
                                    for val in row] 
                      for idx, row in zip(analysis_results['stats_df'].index, analysis_results['stats_df'].values)]
        stats_data.insert(0, ['Metric'] + analysis_results['stats_df'].columns.tolist())
        stats_table = Table(stats_data, colWidths=[1.5*inch, 1.2*inch, 1.2*inch, 1.2*inch])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F2F2F2')])
        ]))
        story.append(stats_table)
        story.append(Spacer(1, 0.3*inch))
        
        # 生成平均光谱图
        story.append(PageBreak())
        story.append(Paragraph("3. Average Spectrum Visualization", styles['Heading2']))
        avg_df = analysis_results['average_spectrum_df']
        
        fig, ax = plt.subplots(figsize=(8, 5))
        ax.plot(avg_df['Wavelength (nm)'], avg_df['Average Intensity'], linewidth=2, color='#1f4788')
        ax.set_title('Average Spectrum', fontsize=14, fontweight='bold')
        ax.set_xlabel('Wavelength (nm)', fontsize=11)
        ax.set_ylabel('Intensity', fontsize=11)
        ax.grid(True, alpha=0.3)
        ax.set_facecolor('#f8f9fa')
        
        img_buffer = BytesIO()
        fig.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)
        plt.close(fig)
        
        img = Image(img_buffer, width=6*inch, height=3.75*inch)
        story.append(img)
        story.append(Spacer(1, 0.3*inch))
        
        # 构建PDF
        doc.build(story)
        print(f"PDF报告已保存至: {pdf_path}")
        return pdf_path
        
    except Exception as e:
        print(f"生成PDF报告时出错: {e}")
        return None


def generate_word_report(report_subfolder, input_filename, analysis_results):
    """
    使用python-docx生成Word格式报告。
    """
    try:
        word_path = os.path.join(report_subfolder, f"{input_filename}_report.docx")
        doc = Document()
        
        # 设置样式
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # 标题
        title = doc.add_heading('Spectrum Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(31, 71, 136)  # 蓝色
        
        # 生成时间
        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.runs[0]
        date_run.font.size = Pt(10)
        date_run.italic = True
        
        doc.add_paragraph()  # 空行
        
        # 1. 峰值指标摘要表
        doc.add_heading('1. Peak Metrics Summary', level=1)
        table1 = doc.add_table(rows=len(analysis_results['summary_df']) + 1, 
                               cols=len(analysis_results['summary_df'].columns))
        table1.style = 'Light Grid Accent 1'
        
        # 表头
        hdr_cells = table1.rows[0].cells
        for i, col_name in enumerate(analysis_results['summary_df'].columns):
            hdr_cells[i].text = str(col_name)
            # 设置表头样式
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            # 设置背景颜色
            shading_elm = parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
            hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
        
        # 数据行
        for row_idx, row_data in enumerate(analysis_results['summary_df'].values, 1):
            cells = table1.rows[row_idx].cells
            for col_idx, value in enumerate(row_data):
                if isinstance(value, float):
                    cells[col_idx].text = f"{value:.4f}"
                else:
                    cells[col_idx].text = str(value)
        
        doc.add_paragraph()  # 空行
        
        # 2. 统计数据表
        doc.add_heading('2. Statistical Summary', level=1)
        table2 = doc.add_table(rows=len(analysis_results['stats_df']) + 1,
                               cols=len(analysis_results['stats_df'].columns) + 1)
        table2.style = 'Light Grid Accent 1'
        
        # 表头
        hdr_cells = table2.rows[0].cells
        hdr_cells[0].text = 'Metric'
        for i, col_name in enumerate(analysis_results['stats_df'].columns, 1):
            hdr_cells[i].text = str(col_name)
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # 数据行
        for row_idx, (idx, row_data) in enumerate(analysis_results['stats_df'].iterrows(), 1):
            cells = table2.rows[row_idx].cells
            cells[0].text = str(idx)
            for col_idx, value in enumerate(row_data, 1):
                cells[col_idx].text = f"{value:.4f}"
        
        doc.add_paragraph()  # 空行
        
        # 3. 平均光谱图表
        doc.add_heading('3. Average Spectrum Visualization', level=1)
        avg_df = analysis_results['average_spectrum_df']
        
        fig, ax = plt.subplots(figsize=(8, 5))
        ax.plot(avg_df['Wavelength (nm)'], avg_df['Average Intensity'], linewidth=2.5, color='#1f4788')
        ax.set_title('Average Spectrum', fontsize=14, fontweight='bold')
        ax.set_xlabel('Wavelength (nm)', fontsize=11)
        ax.set_ylabel('Intensity', fontsize=11)
        ax.grid(True, alpha=0.3)
        ax.set_facecolor('#f8f9fa')
        
        # 保存图表到临时文件
        spectrum_plot_path = os.path.join(report_subfolder, 'spectrum_plot.png')
        fig.savefig(spectrum_plot_path, format='png', dpi=300, bbox_inches='tight')
        plt.close(fig)
        
        # 添加图片到Word文档
        try:
            doc.add_picture(spectrum_plot_path, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as img_err:
            print(f"添加图片到Word文档时出错: {img_err}")
        
        doc.save(word_path)
        print(f"Word报告已保存至: {word_path}")
        
        # 清理临时图表文件
        try:
            if os.path.exists(spectrum_plot_path):
                os.remove(spectrum_plot_path)
        except:
            pass
        
        return word_path
        
    except Exception as e:
        print(f"生成Word报告时出错: {e}")
        return None


def generate_reports(input_path, output_folder, analysis_results, generate_csv, generate_pdf, generate_word):
    """
    根据分析结果和用户选项，生成各种格式的报告文件。
    """
    try:
        # 创建一个唯一的子文件夹来存放本次任务的所有报告
        timestamp = time.strftime("%Y%m%d-%H%M%S")
        input_filename = os.path.splitext(os.path.basename(input_path))[0]
        report_subfolder = os.path.join(output_folder, f"Report_{input_filename}_{timestamp}")
        os.makedirs(report_subfolder, exist_ok=True)

        # --- 生成 Excel 报告 (核心功能) ---
        if generate_csv:  # 虽然选项是CSV，但Excel更强大，能保存多个工作表
            excel_path = os.path.join(report_subfolder, "analysis_summary.xlsx")
            with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
                analysis_results['summary_df'].to_excel(writer, sheet_name='Peak Metrics Summary', index=False)
                analysis_results['stats_df'].to_excel(writer, sheet_name='Statistics')
                analysis_results['average_spectrum_df'].to_excel(writer, sheet_name='Average Spectrum Data',
                                                                 index=False)
                analysis_results['full_spectra_df'].to_excel(writer, sheet_name='All Spectra Data', index=False)
            print(f"Excel 报告已保存至: {excel_path}")

        # --- 生成 PDF 报告 ---
        if generate_pdf:
            generate_pdf_report(report_subfolder, input_filename, analysis_results)

        # --- 生成 Word 报告 ---
        if generate_word:
            generate_word_report(report_subfolder, input_filename, analysis_results)

    except Exception as e:
        print(f"生成报告时发生错误: {e}")